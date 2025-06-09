import os
import docx
import pytesseract
from PIL import Image
import cv2
import numpy as np
import fitz
import tempfile
from pathlib import Path
import tkinter as tk
from tkinter import filedialog

SUPPORTED_LANGS = {
    'Русский': 'rus',
    'English': 'eng',
    'Русский + English': 'rus+eng'
}

def setup_tesseract():
    tesseract_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'
    ]
    
    for path in tesseract_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            os.environ['TESSDATA_PREFIX'] = os.path.dirname(path)
            return True
    
    raise Exception(
        "Tesseract не найден. Пожалуйста, установите Tesseract OCR:\n"
        "1. Скачайте установщик с https://github.com/UB-Mannheim/tesseract/wiki\n"
        "2. Установите Tesseract в папку по умолчанию\n"
        "3. Перезапустите приложение"
    )

def preprocess_image(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    thresh = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
        cv2.THRESH_BINARY, 11, 2
    )
    denoised = cv2.fastNlMeansDenoising(thresh)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    enhanced = clahe.apply(gray)
    return enhanced

def extract_text_from_image(image_path, lang='rus'):
    try:
        image = cv2.imread(image_path)
        if image is None:
            raise Exception(f"Не удалось прочитать изображение: {image_path}")
        
        processed_image = preprocess_image(image)
        
        text = pytesseract.image_to_string(
            processed_image,
            lang=lang,
            config='--psm 3 --oem 3'
        )
        
        data = pytesseract.image_to_data(
            processed_image,
            lang=lang,
            config='--psm 3 --oem 3',
            output_type=pytesseract.Output.DICT
        )
        
        confidences = [int(conf) for conf in data['conf'] if conf != '-1']
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
        
        return text.strip(), avg_confidence
        
    except Exception as e:
        print(f"Ошибка при обработке изображения {image_path}: {str(e)}")
        return "", 0

def extract_text_from_docx(docx_path):
    try:
        doc = docx.Document(docx_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Ошибка при чтении DOCX файла: {str(e)}")
        return ""

def extract_images_from_docx(docx_path, output_folder):
    try:
        doc = docx.Document(docx_path)
        image_count = 0
        image_paths = []
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                image_data = rel.target_part.blob
                image_path = os.path.join(output_folder, f'image_{image_count}.png')
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                image_paths.append(image_path)
        
        return image_paths
    except Exception as e:
        print(f"Ошибка при извлечении изображений из DOCX: {str(e)}")
        return []

def extract_text_from_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print(f"Ошибка при чтении PDF файла: {str(e)}")
        return ""

def extract_images_from_pdf(pdf_path, output_folder):
    try:
        doc = fitz.open(pdf_path)
        image_count = 0
        image_paths = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                image_count += 1
                image_path = os.path.join(output_folder, f'pdf_image_{image_count}.png')
                
                with open(image_path, 'wb') as f:
                    f.write(image_bytes)
                image_paths.append(image_path)
        
        return image_paths
    except Exception as e:
        print(f"Ошибка при извлечении изображений из PDF: {str(e)}")
        return []

def process_document(file_path, output_folder, lang='rus'):
    try:
        os.makedirs(output_folder, exist_ok=True)
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            doc_text = extract_text_from_docx(file_path)
            image_paths = extract_images_from_docx(file_path, output_folder)
        elif file_ext == '.pdf':
            doc_text = extract_text_from_pdf(file_path)
            image_paths = extract_images_from_pdf(file_path, output_folder)
        else:
            raise Exception("Неподдерживаемый формат файла")
        
        results = []
        for img_path in image_paths:
            text, confidence = extract_text_from_image(img_path, lang)
            if text:
                results.append({
                    'filename': os.path.basename(img_path),
                    'text': text,
                    'confidence': confidence
                })
        
        with open(os.path.join(output_folder, 'extracted_text.txt'), 'w', encoding='utf-8') as f:
            f.write("Текст из документа:\n\n")
            f.write(doc_text)
            f.write("\n\nТекст из изображений:\n\n")
            for result in results:
                f.write(f"Изображение: {result['filename']}\n")
                f.write(f"Уверенность: {result['confidence']:.2f}%\n")
                f.write("Текст:\n")
                f.write(result['text'])
                f.write("\n" + "-"*40 + "\n")
        
        return doc_text, results, os.path.join(output_folder, 'extracted_text.txt')
        
    except Exception as e:
        raise Exception(f"Ошибка при обработке документа: {str(e)}")

if __name__ == "__main__":
    # Создаем корневое окно Tkinter (оно будет скрыто)
    root = tk.Tk()
    root.withdraw()

    # Открываем диалог выбора файла
    file_path = filedialog.askopenfilename(
        title="Выберите DOCX файл",
        filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
    )
    
    if not file_path:
        print("Файл не выбран")
        exit()
        
    # Создаем папку для сохранения результатов
    output_folder = "extracted_images"
    
    try:
        doc_text, results, results_file = process_document(file_path, output_folder)
        print(f"Обработка завершена. Результаты сохранены в папку: {output_folder}")
    except Exception as e:
        print(f"Произошла ошибка: {str(e)}") 